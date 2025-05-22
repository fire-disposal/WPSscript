#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
处理Word文档"探索知识海洋.docx"的样式相关操作：
1. 创建"社团01"样式：黑体、小四号、段前段后为1行、单倍行距、大纲级别为1
2. 创建"社团02"样式：宋体、五号、段前段后为0.5行、单倍行距、两端对齐、首行缩进2个字符
3. 应用"社团01"样式到十个大标题（一、二、...十）
4. 应用"社团02"样式到大标题下面的正文
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# 输入和输出文件路径
input_file = "探索知识海洋.docx"
output_file = "探索知识海洋_样式已应用.docx"

def apply_custom_styles():
    """
    创建并应用自定义样式到文档
    """
    # 获取当前脚本所在目录的绝对路径
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # 构建输入和输出文件的绝对路径
    input_path = os.path.join(script_dir, input_file)
    output_path = os.path.join(script_dir, output_file)
    
    print(f"正在处理文件: {input_path}")
    
    try:
        # 打开文档
        doc = Document(input_path)
        
        # 先进行全文清洗
        print("正在进行全文清洗...")
        clean_document(doc)
        
        # 创建并应用自定义样式
        print("正在创建并应用自定义样式...")
        create_and_apply_styles(doc)
        
        # 保存文档
        doc.save(output_path)
        print(f"样式应用完成，文档已保存为: {output_path}")
        
    except Exception as e:
        print(f"处理文档时出错: {e}")
        import traceback
        traceback.print_exc()

def clean_document(doc):
    """
    对文档进行全文清洗：
    1. 删除所有"*"符号（加粗标记）
    2. 清理多余的空行
    
    Args:
        doc: Document对象
    """
    # 删除所有"*"符号
    asterisk_count = 0
    for para in doc.paragraphs:
        if '*' in para.text:
            # 保存原始文本
            original_text = para.text
            # 删除所有"*"符号
            new_text = original_text.replace('*', '')
            # 如果文本有变化，更新段落文本
            if new_text != original_text:
                para.text = new_text
                asterisk_count += original_text.count('*')
    
    print(f"已删除 {asterisk_count} 个'*'符号")
    
    # 清理多余的空行
    empty_count = 0
    consecutive_empty = False
    
    for i, para in enumerate(doc.paragraphs):
        # 如果当前段落为空
        if not para.text.strip():
            # 如果前一个段落也为空，标记为多余空行
            if consecutive_empty:
                para.text = " "  # 设置为空格而不是完全为空，以保持段落的存在
                empty_count += 1
            consecutive_empty = True
        else:
            consecutive_empty = False
    
    print(f"已清理 {empty_count} 个多余的空行")

def create_and_apply_styles(doc):
    """
    创建并应用自定义样式"社团01"和"社团02"
    
    Args:
        doc: Document对象
    """
    # 创建"社团01"样式
    print("正在创建'社团01'样式...")
    style1 = create_style_club01(doc)
    
    # 创建"社团02"样式
    print("正在创建'社团02'样式...")
    style2 = create_style_club02(doc)
    
    # 应用样式到大标题和正文
    print("正在应用样式到文档内容...")
    apply_styles_to_content(doc, style1, style2)
    
    print("样式创建和应用完成")

def create_style_club01(doc):
    """
    创建"社团01"样式：黑体、小四号、段前段后为1行、单倍行距、大纲级别为1
    
    Args:
        doc: Document对象
        
    Returns:
        创建的样式对象
    """
    # 检查样式是否已存在，如果存在则删除
    if "社团01" in doc.styles:
        print("样式'社团01'已存在，将重新创建")
    
    # 创建新样式
    style = doc.styles.add_style("社团01", WD_STYLE_TYPE.PARAGRAPH)
    
    # 设置字体为黑体
    style.font.name = "黑体"
    # 设置中文字体
    style._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
    
    # 设置字号为小四号（12磅）
    style.font.size = Pt(12)
    
    # 设置段前段后为1行（约12磅）
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(12)
    
    # 设置行距为单倍行距
    style.paragraph_format.line_spacing = 1.0
    
    # 设置大纲级别为1
    style.paragraph_format.outline_level = 1
    
    return style

def create_style_club02(doc):
    """
    创建"社团02"样式：宋体、五号、段前段后为0.5行、单倍行距、两端对齐、首行缩进2个字符
    
    Args:
        doc: Document对象
        
    Returns:
        创建的样式对象
    """
    # 检查样式是否已存在，如果存在则删除
    if "社团02" in doc.styles:
        print("样式'社团02'已存在，将重新创建")
    
    # 创建新样式
    style = doc.styles.add_style("社团02", WD_STYLE_TYPE.PARAGRAPH)
    
    # 设置字体为宋体
    style.font.name = "宋体"
    # 设置中文字体
    style._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    
    # 设置字号为五号（10.5磅）
    style.font.size = Pt(10.5)
    
    # 设置段前段后为0.5行（约6磅）
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(6)
    
    # 设置行距为单倍行距
    style.paragraph_format.line_spacing = 1.0
    
    # 设置两端对齐
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # 设置首行缩进2个字符（约21磅，假设一个字符约为10.5磅）
    style.paragraph_format.first_line_indent = Pt(21)
    
    return style

def apply_styles_to_content(doc, style1, style2):
    """
    应用样式到文档内容：
    - "社团01"样式应用到十个大标题（一、二、...十）
    - "社团02"样式应用到大标题下面的正文
    
    Args:
        doc: Document对象
        style1: "社团01"样式对象
        style2: "社团02"样式对象
    """
    # 打印所有段落的文本，帮助调试
    print("\n文档段落内容预览:")
    for i, para in enumerate(doc.paragraphs[:30]):  # 打印前30个段落
        print(f"段落 {i+1}: '{para.text}'")
    
    # 应用样式计数器
    title_count = 0
    content_count = 0
    
    # 定义匹配大标题的函数
    def is_title(text):
        # 移除可能的加粗标记和空白
        clean_text = text.replace('*', '').strip()
        # 匹配"一、"到"十、"的模式，考虑各种标点
        return bool(re.match(r'^[一二三四五六七八九十]+[、.．:：]', clean_text))
    
    # 第一遍：识别并应用样式到所有大标题
    title_paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() and is_title(para.text):
            para.style = "社团01"
            title_count += 1
            title_paragraphs.append(i)
            print(f"已应用'社团01'样式到标题: {para.text}")
    
    # 第二遍：应用样式到大标题之间的正文
    if title_paragraphs:
        for i in range(len(title_paragraphs)):
            start = title_paragraphs[i] + 1
            end = title_paragraphs[i+1] if i+1 < len(title_paragraphs) else len(doc.paragraphs)
            
            for j in range(start, end):
                para = doc.paragraphs[j]
                if para.text.strip() and not is_title(para.text):
                    para.style = "社团02"
                    content_count += 1
                    if content_count <= 5:  # 只打印前5个，避免输出过多
                        print(f"已应用'社团02'样式到正文: '{para.text[:30]}...'")
    
    print(f"共应用'社团01'样式到 {title_count} 个大标题")
    print(f"共应用'社团02'样式到 {content_count} 个正文段落")

if __name__ == "__main__":
    try:
        apply_custom_styles()
    except Exception as e:
        print(f"程序执行出错: {e}")
        import traceback
        traceback.print_exc()